from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi import UploadFile
from pypdf import PdfReader

from .schemas import DocumentType, ExtractedDocument, SourceBlock

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


class DocumentExtractionError(ValueError):
    """Raised when an uploaded document cannot be extracted safely."""


def _normalise_text(value: str) -> str:
    return re.sub(r"[ \t]+", " ", value.replace("\x00", " ")).strip()


def _pdf_blocks(data: bytes, document_id: str) -> list[SourceBlock]:
    reader = PdfReader(BytesIO(data))
    blocks: list[SourceBlock] = []
    for page_index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        page_lines = [
            _normalise_text(line)
            for line in re.split(r"\r?\n", text)
            if _normalise_text(line)
        ]
        for line_index, line in enumerate(page_lines, start=1):
            blocks.append(
                SourceBlock(
                    blockId=f"{document_id}-p{page_index}-b{line_index}",
                    text=line,
                    pageNumber=page_index,
                    paragraphNumber=line_index,
                )
            )
    return blocks


def _docx_blocks(data: bytes, document_id: str) -> list[SourceBlock]:
    document = Document(BytesIO(data))
    blocks: list[SourceBlock] = []
    paragraph_index = 0
    for paragraph in document.paragraphs:
        text = _normalise_text(paragraph.text)
        if not text:
            continue
        paragraph_index += 1
        blocks.append(
            SourceBlock(
                blockId=f"{document_id}-b{paragraph_index}",
                text=text,
                paragraphNumber=paragraph_index,
            )
        )

    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            cells = [_normalise_text(cell.text) for cell in row.cells]
            text = " | ".join(cell for cell in cells if cell)
            if not text:
                continue
            paragraph_index += 1
            blocks.append(
                SourceBlock(
                    blockId=f"{document_id}-table{table_index}-r{row_index}",
                    text=text,
                    paragraphNumber=paragraph_index,
                )
            )
    return blocks


async def extract_upload(
    upload: UploadFile,
    *,
    document_id: str,
    document_type: DocumentType,
    max_bytes: int,
) -> ExtractedDocument:
    filename = upload.filename or "untitled"
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise DocumentExtractionError("PDF와 DOCX 파일만 지원해요.")

    data = await upload.read(max_bytes + 1)
    if len(data) > max_bytes:
        raise DocumentExtractionError(
            f"파일이 너무 커요. {max_bytes // (1024 * 1024)}MB 이하로 업로드해 주세요."
        )
    if not data:
        raise DocumentExtractionError("빈 파일은 분석할 수 없어요.")

    try:
        blocks = (
            _pdf_blocks(data, document_id)
            if extension == ".pdf"
            else _docx_blocks(data, document_id)
        )
    except Exception as exc:  # library errors vary by malformed document type
        raise DocumentExtractionError(
            "파일을 읽지 못했어요. 파일이 손상되지 않았는지 확인해 주세요."
        ) from exc

    warnings: list[str] = []
    if not blocks:
        warnings.append(
            "텍스트를 찾지 못했어요. 스캔 이미지 PDF는 현재 지원하지 않아요."
        )

    return ExtractedDocument(
        documentId=document_id,
        fileName=filename,
        documentType=document_type,
        mimeType=upload.content_type,
        blocks=blocks,
        warnings=warnings,
        characterCount=sum(len(block.text) for block in blocks),
    )
